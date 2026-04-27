from __future__ import annotations

import re
from docx.document import Document


class Preprocessor:
    """Word文档预处理器：清除空白段落、控制字符、多余空白"""

    CONTROL_CHARS_RE = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]")

    def clean(self, doc: Document) -> Document:
        self._remove_blank_paragraphs(doc)
        self._clean_control_characters(doc)
        self._normalize_whitespace(doc)
        return doc

    def _remove_blank_paragraphs(self, doc: Document) -> None:
        A_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"
        M_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
        for para in list(doc.paragraphs):
            if not para.text.strip():
                # 跳过含图片、公式等富内容的段落
                if para._element.find(f".//{{{A_NS}}}blip") is not None:
                    continue
                if para._element.find(f".//{{{M_NS}}}oMath") is not None:
                    continue
                parent = para._element.getparent()
                if parent is not None:
                    parent.remove(para._element)

    def _clean_control_characters(self, doc: Document) -> None:
        for para in doc.paragraphs:
            if self.CONTROL_CHARS_RE.search(para.text):
                for run in para.runs:
                    run.text = self.CONTROL_CHARS_RE.sub("", run.text)

    def _normalize_whitespace(self, doc: Document) -> None:
        """去除每个 run 尾部多余空白。

        仅使用 rstrip 保留 run 间的前导空格，避免跨 run 文字粘连。
        例如 run1="hello " + run2=" world" → rstrip 后 "hello" + " world"。
        """
        for para in doc.paragraphs:
            for run in para.runs:
                run.text = run.text.rstrip()
