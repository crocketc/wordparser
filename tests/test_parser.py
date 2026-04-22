"""测试WordParser主解析器"""

from pathlib import Path
from unittest.mock import MagicMock, Mock, patch

import pytest

from wordparser.config import ParserConfig, TOCPosition
from wordparser.core.models import ParsedDocument
from wordparser.core.parser import WordParser
from wordparser.exceptions import DocumentError, WordParserError


class TestWordParser:
    """测试WordParser类"""

    def test_init_with_default_config(self):
        """测试使用默认配置初始化"""
        parser = WordParser()
        assert parser.config is not None
        assert isinstance(parser.config, ParserConfig)
        assert parser.config.generate_toc is True

    def test_init_with_custom_config(self):
        """测试使用自定义配置初始化"""
        config = ParserConfig(generate_toc=False, max_heading_level=3)
        parser = WordParser(config)
        assert parser.config.generate_toc is False
        assert parser.config.max_heading_level == 3

    def test_init_components(self):
        """测试组件初始化"""
        parser = WordParser()
        assert parser.preprocessor is not None
        assert parser.structure_parser is not None
        assert parser.table_processor is not None
        assert parser.formula_processor is not None
        assert parser.toc_generator is not None
        assert parser.postprocessor is not None

    def test_parse_success(self, tmp_path):
        """测试成功解析文档"""
        # 创建真实的docx文件
        from docx import Document

        docx_file = tmp_path / "test_success.docx"
        doc = Document()
        doc.add_heading("测试文档", level=1)
        doc.add_paragraph("这是一个测试段落。")
        doc.save(str(docx_file))

        # 执行解析
        parser = WordParser()
        result = parser.parse(str(docx_file))

        # 验证返回字符串
        assert isinstance(result, str)
        assert len(result) > 0
        assert "# 测试文档" in result

    def test_parse_nonexistent_file(self):
        """测试解析不存在的文件"""
        parser = WordParser()
        with pytest.raises(DocumentError, match="文件不存在"):
            parser.parse("nonexistent.docx")

    def test_parse_unsupported_format(self, tmp_path):
        """测试解析不支持的文件格式"""
        unsupported_file = tmp_path / "test.txt"
        unsupported_file.write_text("content")

        parser = WordParser()
        with pytest.raises(DocumentError, match="不支持的文件格式"):
            parser.parse(str(unsupported_file))

    def test_parse_with_exception(self, tmp_path):
        """测试解析过程中抛出异常"""
        # 创建一个无效的docx文件
        invalid_file = tmp_path / "invalid.docx"
        invalid_file.write_text("invalid content")

        parser = WordParser()
        with pytest.raises(WordParserError):
            parser.parse(str(invalid_file))

    def test_parse_with_report(self, tmp_path):
        """测试解析并返回报告"""
        # 创建一个真实的docx文件进行测试
        from docx import Document

        docx_file = tmp_path / "test_report.docx"
        doc = Document()
        doc.add_heading("测试文档", level=1)
        doc.add_paragraph("这是一个测试段落。")
        doc.save(str(docx_file))

        # 解析文档
        parser = WordParser()
        markdown, report = parser.parse_with_report(str(docx_file))

        # 验证
        assert isinstance(markdown, str)
        assert len(markdown) > 0
        assert report.success is True
        assert report.stats.total_paragraphs >= 1

    def test_parse_without_toc(self, tmp_path):
        """测试不生成目录的解析"""
        from docx import Document

        docx_file = tmp_path / "test_no_toc.docx"
        doc = Document()
        doc.add_heading("测试文档", level=1)
        doc.add_paragraph("内容")
        doc.save(str(docx_file))

        config = ParserConfig(generate_toc=False)
        parser = WordParser(config)
        result = parser.parse(str(docx_file))

        assert isinstance(result, str)
        assert len(result) > 0

    def test_parse_with_custom_heading_level(self, tmp_path):
        """测试自定义标题级别"""
        from docx import Document

        docx_file = tmp_path / "test_heading.docx"
        doc = Document()
        doc.add_heading("一级标题", level=1)
        doc.add_heading("二级标题", level=2)
        doc.add_heading("三级标题", level=3)
        doc.save(str(docx_file))

        config = ParserConfig(max_heading_level=3)
        parser = WordParser(config)
        result = parser.parse(str(docx_file))

        assert isinstance(result, str)
        assert len(result) > 0
        assert "# 一级标题" in result

    def test_generate_report(self, tmp_path):
        """测试报告生成"""
        # 创建真实的docx文件
        from docx import Document

        docx_file = tmp_path / "test_report_doc.docx"
        doc = Document()
        doc.add_heading("测试文档", level=1)
        doc.add_heading("第一节", level=2)
        doc.add_paragraph("内容1")
        doc.add_paragraph("内容2")
        doc.save(str(docx_file))

        # 执行解析
        parser = WordParser()
        _, report = parser.parse_with_report(str(docx_file))

        # 验证报告
        assert report.success is True
        assert report.stats.total_headings >= 2
        assert report.stats.total_paragraphs >= 2

    def test_toc_position_config(self):
        """测试目录位置配置"""
        config = ParserConfig(toc_position=TOCPosition.BEFORE_TITLE)
        parser = WordParser(config)
        assert parser.config.toc_position == TOCPosition.BEFORE_TITLE

        config = ParserConfig(toc_position=TOCPosition.AFTER_TITLE)
        parser = WordParser(config)
        assert parser.config.toc_position == TOCPosition.AFTER_TITLE
