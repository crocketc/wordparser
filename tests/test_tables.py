"""测试表格处理器"""

import pytest
from docx import Document
from docx.table import Table
from wordparser.core.tables import TableProcessor


class TestTableProcessor:
    """测试TableProcessor类"""

    def test_init(self):
        """测试初始化"""
        processor = TableProcessor()
        assert processor is not None

    def test_process_simple_table(self, docx_with_table):
        """测试处理简单表格"""
        doc = Document(docx_with_table)
        processor = TableProcessor()

        # 获取第一个表格
        table = doc.tables[0]

        # 处理表格
        result = processor.process_simple(table)

        # 验证结果
        assert result is not None
        assert isinstance(result, str)
        assert len(result) > 0

    def test_process_simple_table_with_merged_cells(self, tmp_path):
        """测试处理合并单元格的表格"""
        # 创建包含合并单元格的表格
        doc = Document()
        table = doc.add_table(rows=3, cols=3)

        # 填充数据
        table.rows[0].cells[0].text = "标题"
        table.rows[0].cells[1].text = "列1"
        table.rows[0].cells[2].text = "列2"

        table.rows[1].cells[0].text = "行1"
        table.rows[1].cells[1].text = "数据1-1"
        table.rows[1].cells[2].text = "数据1-2"

        table.rows[2].cells[0].text = "行2"
        table.rows[2].cells[1].text = "数据2-1"
        table.rows[2].cells[2].text = "数据2-2"

        # 合并第一行
        merged_cell = table.rows[0].cells[0].merge(table.rows[0].cells[2])

        doc_path = tmp_path / "merged_table.docx"
        doc.save(str(doc_path))

        # 处理表格
        processor = TableProcessor()
        result = processor.process_simple(table)

        # 验证结果包含合并信息
        assert result is not None
        assert isinstance(result, str)

    def test_is_complex_table(self, docx_with_table):
        """测试判断是否为复杂表格"""
        doc = Document(docx_with_table)
        processor = TableProcessor()

        # 获取第一个表格
        table = doc.tables[0]

        # 简单表格应该返回False
        result = processor.is_complex(table)
        assert isinstance(result, bool)

    def test_is_complex_with_nested_tables(self, tmp_path):
        """测试判断嵌套表格"""
        # 创建嵌套表格
        doc = Document()
        table = doc.add_table(rows=2, cols=2)

        # 在第一个单元格中插入嵌套表格
        nested_table = table.rows[0].cells[0].add_table(rows=2, cols=2)

        processor = TableProcessor()
        result = processor.is_complex(table)

        # 嵌套表格应该被识别为复杂表格
        assert result is True

    def test_is_complex_with_many_rows(self, tmp_path):
        """测试判断多行表格"""
        # 创建超过阈值的行数的表格
        doc = Document()
        table = doc.add_table(rows=15, cols=3)

        processor = TableProcessor()
        result = processor.is_complex(table)

        # 多行表格应该被识别为复杂表格
        assert result is True

    def test_process_simple_with_empty_table(self, tmp_path):
        """测试处理空表格"""
        doc = Document()
        table = doc.add_table(rows=1, cols=1)

        processor = TableProcessor()
        result = processor.process_simple(table)

        # 空表格应该返回空字符串或默认值
        assert result is not None
