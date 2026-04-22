"""测试CLI工具"""

from pathlib import Path
from typer.testing import CliRunner

import pytest

from wordparser_cli.main import app

runner = CliRunner()


class TestCLI:
    """测试CLI命令"""

    def test_parse_nonexistent_file(self):
        """测试解析不存在的文件"""
        result = runner.invoke(app, ["parse", "nonexistent.docx"])
        assert result.exit_code != 0

    def test_parse_simple_document(self, tmp_path):
        """测试解析简单文档"""
        from docx import Document

        # 创建测试文档
        docx_file = tmp_path / "test.docx"
        doc = Document()
        doc.add_heading("测试文档", level=1)
        doc.add_paragraph("这是一个测试段落。")
        doc.save(str(docx_file))

        # 测试解析
        result = runner.invoke(app, ["parse", str(docx_file)])
        assert result.exit_code == 0
        assert "# 测试文档" in result.stdout
        assert "这是一个测试段落。" in result.stdout

    def test_parse_with_output(self, tmp_path):
        """测试解析并输出到文件"""
        from docx import Document

        # 创建测试文档
        docx_file = tmp_path / "test.docx"
        doc = Document()
        doc.add_heading("测试文档", level=1)
        doc.add_paragraph("内容")
        doc.save(str(docx_file))

        # 输出文件
        output_file = tmp_path / "output.md"

        # 测试解析
        result = runner.invoke(app, ["parse", str(docx_file), "-o", str(output_file)])
        assert result.exit_code == 0
        assert output_file.exists()
        content = output_file.read_text(encoding="utf-8")
        assert "# 测试文档" in content

    def test_parse_without_toc(self, tmp_path):
        """测试不生成目录"""
        from docx import Document

        # 创建测试文档
        docx_file = tmp_path / "test.docx"
        doc = Document()
        doc.add_heading("测试文档", level=1)
        doc.add_paragraph("内容")
        doc.save(str(docx_file))

        # 测试解析（不生成目录）
        result = runner.invoke(app, ["parse", str(docx_file), "--no-toc"])
        assert result.exit_code == 0

    def test_parse_with_max_heading(self, tmp_path):
        """测试限制标题级别"""
        from docx import Document

        # 创建测试文档
        docx_file = tmp_path / "test.docx"
        doc = Document()
        doc.add_heading("一级标题", level=1)
        doc.add_heading("二级标题", level=2)
        doc.add_heading("三级标题", level=3)
        doc.save(str(docx_file))

        # 测试解析（限制最大级别为2）
        result = runner.invoke(app, ["parse", str(docx_file), "--max-heading", "2"])
        assert result.exit_code == 0
        assert "# 一级标题" in result.stdout

    def test_parse_verbose(self, tmp_path):
        """测试详细输出模式"""
        from docx import Document

        # 创建测试文档
        docx_file = tmp_path / "test.docx"
        doc = Document()
        doc.add_heading("测试文档", level=1)
        doc.add_paragraph("内容")
        doc.save(str(docx_file))

        # 测试解析（详细模式）
        result = runner.invoke(app, ["parse", str(docx_file), "--verbose"])
        assert result.exit_code == 0
        assert "正在解析文档" in result.stdout
        assert "解析统计" in result.stdout

    def test_version(self):
        """测试版本命令"""
        result = runner.invoke(app, ["version"])
        assert result.exit_code == 0
        assert "WordParser" in result.stdout

    def test_invalid_max_heading(self, tmp_path):
        """测试无效的标题级别"""
        from docx import Document

        # 创建测试文档
        docx_file = tmp_path / "test.docx"
        doc = Document()
        doc.add_heading("测试", level=1)
        doc.save(str(docx_file))

        # 测试无效的标题级别（超过6）
        result = runner.invoke(app, ["parse", str(docx_file), "--max-heading", "10"])
        assert result.exit_code != 0

    def test_help(self):
        """测试帮助信息"""
        result = runner.invoke(app, ["--help"])
        assert result.exit_code == 0
        assert "Word" in result.stdout  # 检查中文字符
        assert "parse" in result.stdout

    def test_parse_command_help(self):
        """测试parse命令帮助"""
        result = runner.invoke(app, ["parse", "--help"])
        assert result.exit_code == 0
        assert "docx_file" in result.stdout
        assert "--output" in result.stdout
        assert "--toc" in result.stdout
