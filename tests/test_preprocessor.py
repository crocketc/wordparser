from docx import Document
from wordparser.core.preprocessor import Preprocessor


def test_remove_blank_paragraphs(docx_with_blank_paragraphs):
    doc = Document(str(docx_with_blank_paragraphs))
    proc = Preprocessor()
    result = proc.clean(doc)
    texts = [p.text for p in result.paragraphs if p.text.strip()]
    assert texts == ["有内容", "也有内容"]


def test_clean_control_characters(tmp_path):
    """通过修改docx ZIP内的XML文件注入控制字符，模拟外部获取的含脏数据文档。
    lxml/Python-docx 不允许 \x00 等控制字符通过API写入，但它们可能存在于
    其他工具生成的文档XML中。这里用 \x7f(DEL) 作为可验证的代表字符。"""
    import zipfile
    from docx import Document
    doc = Document()
    doc.add_paragraph("HelloWorld")  # 占位文本，稍后会被替换为含控制字符的文本
    doc.add_paragraph("正常文本")
    path = tmp_path / "control.docx"
    doc.save(str(path))

    # 先读取所有 ZIP 条目内容
    with zipfile.ZipFile(str(path), "r") as zf:
        entries = {name: zf.read(name) for name in zf.namelist()}

    # 替换为含有 DEL 控制字符的文本
    entries["word/document.xml"] = entries["word/document.xml"].replace(
        b"HelloWorld", b"Hello\x7fWorld"
    )

    # 用修改后的内容重新写入 ZIP
    with zipfile.ZipFile(str(path), "w") as zf:
        for name, data in entries.items():
            zf.writestr(name, data)

    loaded = Document(str(path))
    proc = Preprocessor()
    result = proc.clean(loaded)
    texts = [p.text for p in result.paragraphs]
    assert "\x7f" not in texts[0]
    assert texts[0] == "HelloWorld"


def test_normalize_whitespace(tmp_path):
    from docx import Document
    doc = Document()
    doc.add_paragraph("  多余空格  ")
    doc.add_paragraph("正常")
    path = tmp_path / "whitespace.docx"
    doc.save(str(path))

    loaded = Document(str(path))
    proc = Preprocessor()
    result = proc.clean(loaded)
    text = result.paragraphs[0].text
    assert text == "多余空格"
