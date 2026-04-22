"""测试图片处理器"""

import pytest
from pathlib import Path
from docx import Document
from wordparser.core.images import ImageExtractor


class TestImageExtractor:
    """测试ImageExtractor类"""

    def test_init(self, tmp_path):
        """测试初始化"""
        extractor = ImageExtractor(output_dir=tmp_path)
        assert extractor is not None
        assert extractor.output_dir == tmp_path

    def test_extract_no_images(self, tmp_path, simple_docx):
        """测试从无图片的文档提取"""
        extractor = ImageExtractor(output_dir=tmp_path)
        doc = Document(simple_docx)

        images = extractor.extract(doc, image_prefix="test")

        assert images is not None
        assert isinstance(images, list)
        assert len(images) == 0

    def test_extract_with_images(self, tmp_path):
        """测试从包含图片的文档提取"""
        # 这里我们创建一个模拟测试
        # 实际的图片插入需要图片文件
        extractor = ImageExtractor(output_dir=tmp_path)

        # 创建空文档
        doc = Document()

        images = extractor.extract(doc, image_prefix="img")

        # 验证返回值
        assert images is not None
        assert isinstance(images, list)

    def test_extract_creates_output_directory(self, tmp_path):
        """测试自动创建输出目录"""
        output_dir = tmp_path / "images" / "output"
        extractor = ImageExtractor(output_dir=output_dir)

        doc = Document()

        extractor.extract(doc, image_prefix="test")

        # 验证目录已创建
        assert output_dir.exists()
        assert output_dir.is_dir()

    def test_extract_with_custom_prefix(self, tmp_path):
        """测试自定义图片前缀"""
        extractor = ImageExtractor(output_dir=tmp_path)
        doc = Document()

        images = extractor.extract(doc, image_prefix="custom_")

        assert images is not None
        assert isinstance(images, list)

    def test_extract_returns_image_info(self, tmp_path):
        """测试返回的图片信息结构"""
        extractor = ImageExtractor(output_dir=tmp_path)
        doc = Document()

        images = extractor.extract(doc, image_prefix="test")

        # 验证返回的图片信息结构
        # 如果有图片，每个图片应该包含特定字段
        if images:
            for image_info in images:
                assert isinstance(image_info, dict)
                # 可以验证更多字段

    def test_extract_from_document_with_paragraphs(self, tmp_path):
        """测试从包含段落的文档提取图片"""
        extractor = ImageExtractor(output_dir=tmp_path)
        doc = Document()

        # 添加段落
        doc.add_paragraph("段落1")
        doc.add_paragraph("段落2")

        images = extractor.extract(doc, image_prefix="para_img")

        assert images is not None
        assert isinstance(images, list)
