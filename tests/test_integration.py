"""集成测试

测试 WordParser 的端到端功能，确保所有组件正确协作。
"""

from pathlib import Path
from docx import Document
import pytest

from wordparser import WordParser, ParserConfig, TOCPosition
from wordparser.exceptions import DocumentError, WordParserError


class TestIntegration:
    """集成测试类"""

    @staticmethod
    def _create_rich_docx(tmp_path, filename="test.docx"):
        """创建包含丰富内容的测试文档

        Args:
            tmp_path: 临时目录路径
            filename: 文件名

        Returns:
            文档路径
        """
        docx_file = tmp_path / filename
        doc = Document()

        # 添加标题
        doc.add_heading("第一章：引言", level=1)
        doc.add_heading("1.1 背景", level=2)
        doc.add_paragraph("这是背景介绍段落。")
        doc.add_paragraph("这是第二个段落。")

        doc.add_heading("1.2 目的", level=2)
        doc.add_paragraph("这是目的说明。")

        doc.add_heading("第二章：方法", level=1)
        doc.add_heading("2.1 数据收集", level=2)
        doc.add_paragraph("数据收集的方法说明。")

        # 添加列表
        doc.add_paragraph("方法包括:", style="List Bullet")
        doc.add_paragraph("问卷调查", style="List Bullet")
        doc.add_paragraph("实地访谈", style="List Bullet")

        doc.add_heading("2.2 数据分析", level=2)
        doc.add_paragraph("数据分析方法说明。")

        doc.save(str(docx_file))
        return docx_file

    def test_full_pipeline_with_toc(self, tmp_path):
        """测试完整的解析流程（包含目录）"""
        # 创建测试文档
        docx_file = self._create_rich_docx(tmp_path, "with_toc.docx")

        # 创建解析器（默认生成目录）
        parser = WordParser()

        # 解析文档
        markdown, report = parser.parse_with_report(str(docx_file))

        # 验证结果
        assert isinstance(markdown, str)
        assert len(markdown) > 0
        assert report.success is True

        # 验证标题被正确解析
        assert "# 第一章：引言" in markdown
        assert "## 背景" in markdown
        assert "## 目的" in markdown
        assert "# 第二章：方法" in markdown
        assert "## 数据收集" in markdown
        assert "## 数据分析" in markdown

        # 验证段落被正确解析
        assert "这是背景介绍段落。" in markdown
        assert "这是目的说明。" in markdown

        # 验证列表被正确解析
        assert "-" in markdown or "问卷调查" in markdown

        # 验证统计信息
        assert report.stats.total_headings >= 6  # 至少有6个标题
        assert report.stats.total_paragraphs >= 5  # 至少有5个段落

    def test_full_pipeline_without_toc(self, tmp_path):
        """测试完整的解析流程（不包含目录）"""
        # 创建测试文档
        docx_file = self._create_rich_docx(tmp_path, "without_toc.docx")

        # 创建解析器（不生成目录）
        config = ParserConfig(generate_toc=False)
        parser = WordParser(config)

        # 解析文档
        markdown, report = parser.parse_with_report(str(docx_file))

        # 验证结果
        assert isinstance(markdown, str)
        assert len(markdown) > 0
        assert report.success is True

        # 验证标题被正确解析
        assert "# 第一章：引言" in markdown
        assert "## 背景" in markdown

        # 验证内容完整
        assert "这是背景介绍段落。" in markdown

    def test_nonexistent_file(self):
        """测试解析不存在的文件"""
        parser = WordParser()

        with pytest.raises(DocumentError, match="文件不存在"):
            parser.parse("nonexistent_file.docx")

    def test_invalid_file_format(self, tmp_path):
        """测试解析无效格式的文件"""
        # 创建一个非 docx 文件
        invalid_file = tmp_path / "invalid.txt"
        invalid_file.write_text("这是一个文本文件，不是 Word 文档。")

        parser = WordParser()

        with pytest.raises(DocumentError, match="不支持的文件格式"):
            parser.parse(str(invalid_file))

    def test_corrupted_file(self, tmp_path):
        """测试解析损坏的文件"""
        # 创建一个损坏的 docx 文件
        corrupted_file = tmp_path / "corrupted.docx"
        corrupted_file.write_text("这不是有效的 docx 文件内容")

        parser = WordParser()

        with pytest.raises(WordParserError):
            parser.parse(str(corrupted_file))

    def test_empty_document(self, tmp_path):
        """测试解析空文档"""
        # 创建空文档
        docx_file = tmp_path / "empty.docx"
        doc = Document()
        doc.save(str(docx_file))

        parser = WordParser()
        markdown, report = parser.parse_with_report(str(docx_file))

        # 应该成功解析，但内容为空或很少
        assert report.success is True
        assert isinstance(markdown, str)

    def test_document_with_only_headings(self, tmp_path):
        """测试只有标题的文档"""
        docx_file = tmp_path / "only_headings.docx"
        doc = Document()

        doc.add_heading("标题1", level=1)
        doc.add_heading("标题2", level=2)
        doc.add_heading("标题3", level=3)

        doc.save(str(docx_file))

        parser = WordParser()
        markdown, report = parser.parse_with_report(str(docx_file))

        assert report.success is True
        assert "# 标题1" in markdown
        assert "## 标题2" in markdown
        assert "### 标题3" in markdown
        assert report.stats.total_headings == 3

    def test_document_with_only_paragraphs(self, tmp_path):
        """测试只有段落的文档"""
        docx_file = tmp_path / "only_paragraphs.docx"
        doc = Document()

        doc.add_paragraph("段落1")
        doc.add_paragraph("段落2")
        doc.add_paragraph("段落3")

        doc.save(str(docx_file))

        parser = WordParser()
        markdown, report = parser.parse_with_report(str(docx_file))

        assert report.success is True
        assert "段落1" in markdown
        assert "段落2" in markdown
        assert "段落3" in markdown
        assert report.stats.total_paragraphs >= 3

    def test_document_with_nested_headings(self, tmp_path):
        """测试嵌套标题的文档"""
        docx_file = tmp_path / "nested_headings.docx"
        doc = Document()

        doc.add_heading("一级标题", level=1)
        doc.add_paragraph("一级内容")
        doc.add_heading("二级标题A", level=2)
        doc.add_paragraph("二级内容A")
        doc.add_heading("三级标题", level=3)
        doc.add_paragraph("三级内容")
        doc.add_heading("二级标题B", level=2)
        doc.add_paragraph("二级内容B")

        doc.save(str(docx_file))

        parser = WordParser()
        markdown, report = parser.parse_with_report(str(docx_file))

        assert report.success is True
        assert "# 一级标题" in markdown
        assert "## 二级标题A" in markdown
        assert "### 三级标题" in markdown
        assert "## 二级标题B" in markdown

    def test_chinese_characters(self, tmp_path):
        """测试中文字符处理"""
        docx_file = tmp_path / "chinese.docx"
        doc = Document()

        doc.add_heading("中文标题", level=1)
        doc.add_paragraph("这是中文段落。")
        doc.add_paragraph("包含特殊字符：!@#$%^&*()")

        doc.save(str(docx_file))

        parser = WordParser()
        markdown, report = parser.parse_with_report(str(docx_file))

        assert report.success is True
        assert "中文标题" in markdown
        assert "这是中文段落。" in markdown
        assert "!@#$%^&*()" in markdown

    def test_save_to_file(self, tmp_path):
        """测试保存到文件"""
        # 创建测试文档
        docx_file = self._create_rich_docx(tmp_path, "save_test.docx")

        # 解析文档
        parser = WordParser()
        markdown, report = parser.parse_with_report(str(docx_file))

        # 保存到文件
        output_file = tmp_path / "output.md"
        output_file.write_text(markdown, encoding="utf-8")

        # 验证文件
        assert output_file.exists()
        content = output_file.read_text(encoding="utf-8")
        assert content == markdown
        assert "# 第一章：引言" in content

    def test_custom_max_heading_level(self, tmp_path):
        """测试自定义最大标题级别"""
        docx_file = tmp_path / "heading_level.docx"
        doc = Document()

        doc.add_heading("一级", level=1)
        doc.add_heading("二级", level=2)
        doc.add_heading("三级", level=3)
        doc.add_heading("四级", level=4)
        doc.add_heading("五级", level=5)
        doc.add_heading("六级", level=6)

        doc.save(str(docx_file))

        # 限制最大级别为3
        config = ParserConfig(max_heading_level=3)
        parser = WordParser(config)
        markdown, report = parser.parse_with_report(str(docx_file))

        assert report.success is True
        assert "# 一级" in markdown
        assert "## 二级" in markdown
        assert "### 三级" in markdown

    def test_multiple_documents(self, tmp_path):
        """测试批量处理多个文档"""
        # 创建多个文档
        files = []
        for i in range(3):
            docx_file = tmp_path / f"doc{i}.docx"
            doc = Document()
            doc.add_heading(f"文档{i}", level=1)
            doc.add_paragraph(f"内容{i}")
            doc.save(str(docx_file))
            files.append(docx_file)

        # 批量处理
        parser = WordParser()
        results = []

        for docx_file in files:
            markdown, report = parser.parse_with_report(str(docx_file))
            results.append((docx_file.name, markdown, report))

        # 验证结果
        assert len(results) == 3
        for filename, markdown, report in results:
            assert report.success is True
            assert len(markdown) > 0
