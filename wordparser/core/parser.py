"""WordParser主解析器"""
from __future__ import annotations

import logging
import re
from concurrent.futures import ThreadPoolExecutor, as_completed
from pathlib import Path
from typing import Any

from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from wordparser.config import ParserConfig
from wordparser.core.chart_extractor import ChartExtractor
from wordparser.core.formulas import FormulaProcessor
from wordparser.core.models import ParsedDocument, TitleNode
from wordparser.core.postprocess import PostProcessor
from wordparser.core.preprocessor import Preprocessor
from wordparser.core.renderer import DocumentRenderer
from wordparser.core.report import ParseReport
from wordparser.core.smartart_extractor import SmartArtExtractor
from wordparser.core.structure import StructureParser
from wordparser.core.tables import TableProcessor
from wordparser.core.toc import TOCGenerator, Heading
from wordparser.exceptions import DocumentError, WordParserError
from wordparser.multimodal.parser import MultimodalParser

logger = logging.getLogger(__name__)

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
M_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
A_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"
R_URI = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"

# 占位符格式：不会出现在正常文本中
_PH_IMG = "[__WP_IMG_{}__]"
_PH_CHART = "[__WP_CHART_{}__]"
_PH_SA = "[__WP_SA_{}__]"
_PH_TABLE = "[__WP_TABLE_{}__]"


class WordParser:
    """Word文档转Markdown解析器"""

    def __init__(self, config: ParserConfig | None = None) -> None:
        self.config = config or ParserConfig()
        self._init_components()

    def _init_components(self) -> None:
        self.preprocessor = Preprocessor()
        self.structure_parser = StructureParser(self.config)
        self.table_processor = TableProcessor()
        self.formula_processor = FormulaProcessor()
        self.toc_generator = TOCGenerator()
        self.postprocessor = PostProcessor()

        self.chart_extractor = ChartExtractor()
        self.smartart_extractor = SmartArtExtractor()

        self.renderer = DocumentRenderer(
            libreoffice_path=self.config.libreoffice_path,
        )

        self.vision_client = None
        if self.config.multimodal:
            from wordparser.multimodal.client import OpenAICompatibleVisionClient
            self.vision_client = OpenAICompatibleVisionClient(
                base_url=self.config.multimodal.model.base_url,
                model=self.config.multimodal.model.model,
                temperature=self.config.multimodal.model.temperature,
            )

        self.multimodal_parser = MultimodalParser(
            vision_client=self.vision_client,
            renderer=self.renderer,
            enable_render_fallback=self.config.enable_render_fallback,
        )

    def parse(self, docx_path: str | Path) -> str:
        document = self._parse_document(docx_path)
        return document.metadata.get("markdown", "")

    def parse_with_report(self, docx_path: str | Path) -> tuple[str, ParseReport]:
        document = self._parse_document(docx_path)
        report = self._generate_report(document)
        markdown = document.metadata.get("markdown", "")
        return markdown, report

    def _ensure_docx(self, input_path: Path) -> Path:
        if self.renderer.is_doc(input_path):
            if not self.renderer.is_available():
                raise DocumentError(
                    "解析 .doc 文件需要 LibreOffice，请安装 LibreOffice 或配置 libreoffice_path"
                )
            logger.info(f"检测到 .doc 文件，自动转换为 .docx: {input_path}")
            return self.renderer.convert_doc_to_docx(input_path)
        return input_path

    # ================================================================
    # 主解析流程
    # ================================================================

    def _parse_document(self, docx_path: str | Path) -> ParsedDocument:
        docx_path = Path(docx_path)

        if not docx_path.exists():
            raise DocumentError(f"文件不存在: {docx_path}")

        docx_path = self._ensure_docx(docx_path)

        supported = {".docx", ".doc"}
        if docx_path.suffix.lower() not in supported:
            raise DocumentError(f"不支持的文件格式: {docx_path.suffix}")

        try:
            # 1. 加载文档
            from docx import Document as DocxDocument
            doc = DocxDocument(str(docx_path))

            # 2. 预处理前：注入占位符（图片/Chart/SmartArt → 文本标记）
            rich_ids = self._inject_placeholders(doc)

            # 3. 预处理（占位符作为文本保留，不会被删除）
            doc = self.preprocessor.clean(doc)

            # 4. 设置 doc_part 用于超链接解析
            self.structure_parser.set_doc_part(doc.part)

            # 5. 结构解析（占位符文本会被保留在 paragraph content 中）
            blocks = self.structure_parser.parse(doc)
            title_tree = self.structure_parser.get_title_tree()

            # 6. 并行解析富内容
            image_descs, chart_descs, sa_descs, table_descs = self._parse_rich_content(
                doc, docx_path, rich_ids,
            )

            # 7. 构建 Markdown
            markdown = self._build_markdown(doc, blocks, title_tree)

            # 8. 回填占位符
            markdown = self._resolve_placeholders(markdown, image_descs, chart_descs, sa_descs, table_descs)

            # 9. 后处理
            markdown = self.postprocessor.process(markdown)

            # 10. TOC
            if self.config.generate_toc and title_tree:
                toc_md = self._generate_toc_markdown(title_tree)
                markdown = f"## 目录\n\n{toc_md}\n\n---\n\n{markdown}"

            # 11. 可选：页眉页脚
            if self.config.include_header_footer:
                hf_md = self._extract_header_footer(doc)
                if hf_md:
                    markdown += f"\n\n## 页眉页脚\n\n{hf_md}"

            document = ParsedDocument(
                metadata={
                    "docx_path": str(docx_path),
                    "word_count": sum(len(p.text) for p in doc.paragraphs),
                    "paragraph_count": len(doc.paragraphs),
                    "image_count": len(image_descs),
                    "chart_count": len(chart_descs),
                    "smartart_count": len(sa_descs),
                    "complex_table_count": 0,
                    "markdown": markdown,
                },
                title_tree=title_tree,
                content_blocks=blocks,
            )

            return document

        except Exception as e:
            if isinstance(e, (DocumentError, WordParserError)):
                raise
            raise WordParserError(f"文档解析失败: {e}") from e

    # ================================================================
    # 占位符注入（预处理前）
    # ================================================================

    def _inject_placeholders(self, doc) -> dict:
        """预处理前：在含图片/Chart/SmartArt 的段落中注入文本占位符。

        占位符作为普通文本 run 加入段落，确保预处理不会删除这些段落。
        返回 {type: {id: rId_or_index}} 映射。
        """
        rich_ids = {"img": {}, "chart": {}, "sa": {}, "table": {}}
        body = doc.element.body
        chart_idx = 0
        sa_idx = 0
        table_idx = 0

        # 收集复杂表格（用于后续并行解析）
        # 不在这里插入占位符，而是在 _build_markdown 时处理

        for para_elem in body.iter(f"{{{W_NS}}}p"):
            # 图片
            for blip in para_elem.iter(f"{{{A_NS}}}blip"):
                rId = blip.get(f"{{{R_URI}}}embed")
                if rId:
                    self._append_text_run(para_elem, _PH_IMG.format(rId))
                    rich_ids["img"][rId] = rId

            # Chart
            for chart_ref in para_elem.iter(f"{{{W_NS}}}object"):
                # Chart 通常在 w:object 中
                pass
            # 也检查 w:drawing 中的 chart
            for elem in para_elem.iter():
                local = elem.tag.split("}")[1] if "}" in elem.tag else elem.tag
                if local == "chart":
                    rId = elem.get(f"{{{R_URI}}}id")
                    if rId:
                        ph = _PH_CHART.format(chart_idx)
                        self._append_text_run(para_elem, ph)
                        rich_ids["chart"][chart_idx] = rId
                        chart_idx += 1

            # SmartArt (diagramData / dgm)
            for elem in para_elem.iter():
                local = elem.tag.split("}")[1] if "}" in elem.tag else elem.tag
                if "dgm" in elem.tag.lower() or "diagram" in elem.tag.lower():
                    ph = _PH_SA.format(sa_idx)
                    self._append_text_run(para_elem, ph)
                    rich_ids["sa"][sa_idx] = sa_idx
                    sa_idx += 1
                    break

        return rich_ids

    @staticmethod
    def _append_text_run(para_elem, text: str) -> None:
        """在段落末尾追加一个纯文本 run"""
        run = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.text = text
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        run.append(t)
        para_elem.append(run)

    # ================================================================
    # 富内容并行解析
    # ================================================================

    def _parse_rich_content(self, doc, docx_path: Path, rich_ids: dict):
        """并行解析图片、Chart、SmartArt、复杂表格"""
        max_workers = self.config.multimodal.max_concurrent if self.config.multimodal else 2

        image_descs = self._parse_images_parallel(doc, rich_ids.get("img", {}), max_workers)
        chart_descs = self._parse_charts(docx_path)
        sa_descs = self._parse_smartarts(docx_path)
        table_descs = self._parse_complex_tables_parallel(doc, rich_ids.get("table", {}), max_workers)

        return image_descs, chart_descs, sa_descs, table_descs

    def _parse_images_parallel(self, doc, img_ids: dict, max_workers: int) -> dict[str, str]:
        """并行解析嵌入图片"""
        if not img_ids:
            return {}

        self._ensure_vision_client()

        from wordparser.multimodal.prompts import IMAGE_PROMPT

        def _parse_one(rId):
            try:
                rel = doc.part.rels[rId]
                image_bytes = rel.target_part.blob
                desc = self.vision_client.parse_from_bytes(image_bytes, IMAGE_PROMPT)
                return rId, desc
            except Exception as e:
                return rId, f"[图片解析失败: {e}]"

        image_map: dict[str, str] = {}
        with ThreadPoolExecutor(max_workers=max_workers) as pool:
            futures = {pool.submit(_parse_one, rId): rId for rId in img_ids}
            for future in as_completed(futures):
                rId, desc = future.result()
                image_map[rId] = desc

        return image_map

    def _parse_complex_tables_parallel(self, doc, table_ids: dict, max_workers: int) -> dict[int, str]:
        """并行解析复杂表格，返回 {table_element_id: description}"""
        if not table_ids:
            return {}

        self._ensure_vision_client()

        # 收集所有复杂表格
        complex_tables = []
        body = doc.element.body
        for tbl_elem in body.iter(f"{{{W_NS}}}tbl"):
            table = None
            for t in doc.tables:
                if t._element is tbl_elem:
                    table = t
                    break
            if table and self.table_processor.is_complex(table):
                complex_tables.append((tbl_elem, table))

        if not complex_tables:
            return {}

        def _parse_one(tbl_elem, table):
            try:
                table_data = self.table_processor.extract_table_data(table)
                result = self.multimodal_parser.parse_complex_table(table_data)
                return tbl_elem, result.content
            except Exception as e:
                return tbl_elem, f"[复杂表格解析失败: {e}]"

        table_map: dict[int, str] = {}
        with ThreadPoolExecutor(max_workers=max_workers) as pool:
            futures = {pool.submit(_parse_one, tbl_elem, table): tbl_elem for tbl_elem, table in complex_tables}
            for future in as_completed(futures):
                tbl_elem, desc = future.result()
                table_map[id(tbl_elem)] = desc

        return table_map

    def _ensure_vision_client(self) -> None:
        """延迟初始化 vision_client，从配置读取参数"""
        if not self.vision_client:
            from wordparser.multimodal.client import OpenAICompatibleVisionClient
            # 从配置读取模型参数，默认使用 config.py 中的默认值
            model_config = self.config.multimodal.model if self.config.multimodal else None
            self.vision_client = OpenAICompatibleVisionClient(
                base_url=model_config.base_url if model_config else "http://localhost:1234/v1",
                model=model_config.model if model_config else "qwen3.5-4b",
                temperature=model_config.temperature if model_config else 0.0,
            )
            self.multimodal_parser.vision_client = self.vision_client

    def _wrap_multimodal_result(self, content: str, content_type: str) -> str:
        """将多模态解析结果用代码框包裹

        Args:
            content: AI 解析返回的内容
            content_type: 内容类型（image/chart/smartart/table）

        Returns:
            包裹后的 Markdown 字符串
        """
        return f"\n```{content_type}\n{content}\n```\n"

    def _parse_charts(self, docx_path: Path) -> dict[int, str]:
        """并行解析 Chart 图表，返回 {index: description}"""
        if not self.vision_client:
            return {}

        try:
            chart_data_list = self.chart_extractor.extract(docx_path)
        except Exception as e:
            logger.warning(f"Chart 提取失败: {e}")
            return {}

        if not chart_data_list:
            return {}

        max_workers = self.config.multimodal.max_concurrent if self.config.multimodal else 2

        def _parse_one(index: int, chart_data):
            try:
                r = self.multimodal_parser.parse_chart_with_data(chart_data, docx_path)
                return index, r.content
            except Exception as e:
                return index, f"[图表解析失败: {e}]"

        result: dict[int, str] = {}
        with ThreadPoolExecutor(max_workers=max_workers) as pool:
            futures = {pool.submit(_parse_one, i, chart_data): i for i, chart_data in enumerate(chart_data_list)}
            for future in as_completed(futures):
                index, desc = future.result()
                result[index] = desc
        return result

    def _parse_smartarts(self, docx_path: Path) -> dict[int, str]:
        """并行解析 SmartArt，返回 {index: description}"""
        if not self.vision_client:
            return {}

        try:
            sa_data_list = self.smartart_extractor.extract(docx_path)
        except Exception as e:
            logger.warning(f"SmartArt 提取失败: {e}")
            return {}

        if not sa_data_list:
            return {}

        max_workers = self.config.multimodal.max_concurrent if self.config.multimodal else 2

        def _parse_one(index: int, sa_data):
            try:
                r = self.multimodal_parser.parse_smartart_with_data(sa_data, docx_path)
                return index, r.content
            except Exception as e:
                return index, f"[SmartArt解析失败: {e}]"

        result: dict[int, str] = {}
        with ThreadPoolExecutor(max_workers=max_workers) as pool:
            futures = {pool.submit(_parse_one, i, sa_data): i for i, sa_data in enumerate(sa_data_list)}
            for future in as_completed(futures):
                index, desc = future.result()
                result[index] = desc
        return result

    # ================================================================
    # Markdown 构建
    # ================================================================

    def _build_markdown(self, doc, blocks: list, title_tree: list[TitleNode]) -> str:
        """构建 Markdown，保持 w:p / w:tbl 文档流顺序。"""
        from wordparser.core.models import BlockType

        para_block_map: dict = {}
        for block in blocks:
            pe = block.metadata.get("_para_element")
            if pe is not None:
                para_block_map[id(pe)] = block

        body = doc.element.body
        markdown_lines: list[str] = []
        emitted_blocks: set[int] = set()

        for child in body:
            tag = child.tag

            if tag == f"{{{W_NS}}}p":
                child_id = id(child)
                block = para_block_map.get(child_id)
                if block is not None:
                    self._append_block_markdown(block, markdown_lines)
                    emitted_blocks.add(id(block))

            elif tag == f"{{{W_NS}}}tbl":
                self._append_table_markdown(doc, child, markdown_lines)

        for block in blocks:
            if id(block) not in emitted_blocks:
                self._append_block_markdown(block, markdown_lines)

        return "\n".join(markdown_lines)

    def _resolve_placeholders(
        self,
        markdown: str,
        image_descs: dict[str, str],
        chart_descs: dict[int, str],
        sa_descs: dict[int, str],
        table_descs: dict[int, str] = None,
    ) -> str:
        """将占位符替换为实际的富内容描述"""
        # 图片
        for rId, desc in image_descs.items():
            ph = _PH_IMG.format(rId)
            if ph in markdown:
                markdown = markdown.replace(ph, self._wrap_multimodal_result(desc, "image"))

        # Chart
        for idx, desc in chart_descs.items():
            ph = _PH_CHART.format(idx)
            if ph in markdown:
                markdown = markdown.replace(ph, self._wrap_multimodal_result(desc, "chart"))

        # SmartArt
        for idx, desc in sa_descs.items():
            ph = _PH_SA.format(idx)
            if ph in markdown:
                markdown = markdown.replace(ph, self._wrap_multimodal_result(desc, "smartart"))

        # 复杂表格
        if table_descs:
            for tbl_elem_id, desc in table_descs.items():
                ph = _PH_TABLE.format(tbl_elem_id)
                if ph in markdown:
                    markdown = markdown.replace(ph, self._wrap_multimodal_result(desc, "table"))

        return markdown

    def _append_block_markdown(self, block, lines: list[str]) -> None:
        from wordparser.core.models import BlockType

        if block.type == BlockType.HEADING:
            node = block.content
            lines.append(f"{'#' * node.level} {node.text}\n")
        elif block.type == BlockType.PARAGRAPH:
            lines.append(f"{block.content}\n")
        elif block.type == BlockType.LIST:
            level = block.metadata.get("level", 0)
            indent = "  " * level
            lines.append(f"{indent}- {block.content}\n")

    def _append_table_markdown(self, doc, tbl_element, lines: list[str]) -> None:
        table = None
        for t in doc.tables:
            if t._element is tbl_element:
                table = t
                break

        if table is None:
            return

        if self.table_processor.is_complex(table):
            # 注入占位符，在 _resolve_placeholders 中替换
            tbl_elem_id = id(tbl_element)
            lines.append(_PH_TABLE.format(tbl_elem_id))
        else:
            md = self._simple_table_to_md(table)
            lines.append(f"\n{md}\n")

    def _simple_table_to_md(self, table) -> str:
        if not table or not table.rows:
            return ""

        lines = []
        for row_idx, row in enumerate(table.rows):
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            lines.append(f"| {' | '.join(cells)} |")
            if row_idx == 0:
                col_count = len(row.cells)
                lines.append(f"| {' | '.join(['---'] * col_count)} |")

        return "\n".join(lines)

    # ================================================================
    # TOC / Report / Header-Footer
    # ================================================================

    def _generate_toc_markdown(self, title_tree: list[TitleNode]) -> str:
        headings = []
        self._collect_headings(title_tree, headings)
        if not headings:
            return ""
        return self.toc_generator.generate(headings, add_anchors=True)

    def _collect_headings(self, nodes: list[TitleNode], headings: list[Heading], prefix: str = "") -> None:
        counters: dict[int, int] = {}
        for node in nodes:
            level = node.level
            counters[level] = counters.get(level, 0) + 1
            for k in list(counters.keys()):
                if k > level:
                    del counters[k]
            parts = [str(counters[l]) for l in sorted(counters.keys())]
            number = ".".join(parts)
            headings.append(Heading(level=node.level, text=node.text, number=number))
            self._collect_headings(node.children, headings, number)

    def _generate_report(self, document: ParsedDocument) -> ParseReport:
        from wordparser.core.report import ParseStats
        metadata = document.metadata

        def count_titles(nodes):
            return sum(1 + count_titles(n.children) for n in nodes)

        stats = ParseStats(
            total_headings=count_titles(document.title_tree),
            total_paragraphs=metadata.get("paragraph_count", 0),
            total_tables=metadata.get("complex_table_count", 0),
            total_images=metadata.get("image_count", 0),
            multimodal_calls=0,
            multimodal_failures=0,
            processing_time=0.0,
        )
        return ParseReport(success=True, output_path=None, errors=[], stats=stats)

    def _extract_header_footer(self, doc) -> str:
        parts = []
        for section in doc.sections:
            header = section.header
            if header and not header.is_linked_to_previous:
                for p in header.paragraphs:
                    t = p.text.strip()
                    if t:
                        parts.append(f"**页眉**: {t}")
            footer = section.footer
            if footer and not footer.is_linked_to_previous:
                for p in footer.paragraphs:
                    t = p.text.strip()
                    if t:
                        parts.append(f"**页脚**: {t}")
        return "\n".join(parts)
