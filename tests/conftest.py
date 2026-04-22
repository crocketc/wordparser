"""共享测试fixtures"""
import pytest
from pathlib import Path

FIXTURES_DIR = Path(__file__).parent / "fixtures" / "samples"

@pytest.fixture(scope="session")
def fixtures_dir():
    return FIXTURES_DIR


from docx import Document
from docx.shared import Pt


@pytest.fixture
def simple_docx(tmp_path):
    """创建一个简单的测试用Word文档"""
    doc = Document()
    doc.add_heading("测试标题", level=1)
    doc.add_paragraph("这是第一个段落。")
    doc.add_paragraph("这是第二个段落。")
    path = tmp_path / "simple.docx"
    doc.save(str(path))
    return path


@pytest.fixture
def docx_with_blank_paragraphs(tmp_path):
    """创建含空白段落的Word文档"""
    doc = Document()
    doc.add_paragraph("有内容")
    doc.add_paragraph("")
    doc.add_paragraph("   ")
    doc.add_paragraph("\t")
    doc.add_paragraph("也有内容")
    path = tmp_path / "blanks.docx"
    doc.save(str(path))
    return path


@pytest.fixture
def docx_with_headings(tmp_path):
    """多级标题文档"""
    doc = Document()
    doc.add_heading("一级标题", level=1)
    doc.add_paragraph("正文段落1")
    doc.add_heading("二级标题", level=2)
    doc.add_paragraph("正文段落2")
    doc.add_heading("三级标题", level=3)
    doc.add_paragraph("正文段落3")
    path = tmp_path / "headings.docx"
    doc.save(str(path))
    return path


@pytest.fixture
def docx_with_numbered_headings(tmp_path):
    """带序号的标题"""
    doc = Document()
    doc.add_heading("1. 项目背景", level=1)
    doc.add_heading("1.1 技术方案", level=2)
    doc.add_heading("1.1.1 详细设计", level=3)
    path = tmp_path / "numbered.docx"
    doc.save(str(path))
    return path
