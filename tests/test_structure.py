from docx import Document
from wordparser.core.structure import StructureParser
from wordparser.core.models import BlockType


def test_parse_headings(docx_with_headings):
    doc = Document(str(docx_with_headings))
    parser = StructureParser()
    blocks = parser.parse(doc)

    headings = [b for b in blocks if b.type == BlockType.HEADING]
    assert len(headings) == 3
    assert headings[0].content.level == 1
    assert headings[0].content.text == "一级标题"
    assert headings[1].content.level == 2
    assert headings[2].content.level == 3


def test_parse_paragraphs(docx_with_headings):
    doc = Document(str(docx_with_headings))
    parser = StructureParser()
    blocks = parser.parse(doc)

    paras = [b for b in blocks if b.type == BlockType.PARAGRAPH]
    assert len(paras) == 3
    assert paras[0].content == "正文段落1"


def test_strip_heading_numbers(docx_with_numbered_headings):
    doc = Document(str(docx_with_numbered_headings))
    parser = StructureParser()
    blocks = parser.parse(doc)

    headings = [b for b in blocks if b.type == BlockType.HEADING]
    assert headings[0].content.text == "项目背景"
    assert headings[1].content.text == "技术方案"
    assert headings[2].content.text == "详细设计"


def test_title_tree_built(docx_with_headings):
    doc = Document(str(docx_with_headings))
    parser = StructureParser()
    blocks = parser.parse(doc)

    tree = parser.get_title_tree()
    assert len(tree) == 1
    assert tree[0].level == 1
    assert tree[0].text == "一级标题"
    assert len(tree[0].children) == 1
    assert tree[0].children[0].level == 2


def test_heading_anchor_generation(docx_with_headings):
    doc = Document(str(docx_with_headings))
    parser = StructureParser()
    blocks = parser.parse(doc)

    headings = [b for b in blocks if b.type == BlockType.HEADING]
    assert headings[0].content.anchor == "一级标题"
    assert headings[1].content.anchor == "二级标题"
